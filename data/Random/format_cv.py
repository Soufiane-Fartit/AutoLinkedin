import re
from docx import Document

"""
def string2xml(xmlstring):
    import xml.etree.ElementTree as ET
    tree = ET.ElementTree(ET.fromstring(xmlstring))
    return tree

def string2docxml(xmlstring):
    from docx.oxml.xmlchemy import XmlString
    return XmlString(xmlstring)

def xml2string(xml):
    return str(xml)

def replaceInXml(xml, string1, string2):
    s = xml2string(xml)
    s = re.sub(string1, string2, s)
    x = string2docxml(s)
    return x

def replaceInDoc(doc, string1, string2):
    xml_ = doc.element.xml
    xml_ = replaceInXml(xml_, string1, string2)
    doc.element.xml = xml_
    return doc
"""

def replaceInDoc(doc, string1, string2):
    for paragraph in doc.paragraphs:
        if string1 in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if string1 in inline[i].text:
                    text = inline[i].text.replace(string1, string2)
                    inline[i].text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if string1 in paragraph.text:
                        inline = paragraph.runs
                        for i in range(len(inline)):
                            if string1 in inline[i].text:
                                text = inline[i].text.replace(string1, string2)
                                inline[i].text = text
    return doc

def quickReplaceInDoc(doc, string1, string2):
    for paragraph in doc.paragraphs:
        if string1 in paragraph.text:
            paragraph.text = paragraph.text.replace(string1, string2)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if string1 in paragraph.text:
                        paragraph.text = paragraph.text.replace(string1, string2)

    return doc

def replaceDictInDoc(doc, dic):
    for k,v in dic.items():
        doc = replaceInDoc(doc, k, v)
    return doc

def quickReplaceDictInDoc(doc, dic):
    for k,v in dic.items():
        doc = quickReplaceInDoc(doc, k, v)
    return doc

def addSkills(doc, skills):
    for paragraph in doc.paragraphs:
        if '<SKILLS>' in paragraph.text:
            paragraph.text = ""
            p = paragraph._p
            for skill in skills:
                para = doc.add_paragraph(skill, style='List Bullet')
                p.addnext(para._p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '<SKILLS>' in paragraph.text:
                        paragraph.text = ""
                        p = paragraph._p
                        for skill in skills:
                            para = doc.add_paragraph(skill, style='List Bullet')
                            p.addnext(para._p)

    return doc

def addPicture(doc, img_path):
    from docx.shared import Inches

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "<PICTURE>" in paragraph.text:
                        cell._element.clear_content()
                        cell.add_paragraph().add_run().add_picture(img_path, width=Inches(1.8))
    return doc

doc = Document("template_01.docx")
doc = replaceDictInDoc(doc, {'<PRENOM>':'SOUFIANE', 
                             '<NOM>':'FARTIT',
                             '<EXP1>':'Ingénieur en Machine Learning',
                             '<COMPANY1>':'FM Logistic',
                             '<EXP1CITY>':'Phalsbourg'})
doc = quickReplaceDictInDoc(doc, {'<ADDRESS>':'strasbourg, 67000, France',
                             '<TEL>':'0662833264',
                             '<EMAIL>':'soufiane.fartit@gmail.com',
                             '<PROFIL>':'je suis motivé pour faire ça et ça dans le secteur ect... je possède un diplome dans ça et ça... blabla',
                             '<LANGUAGE1>':'ANGLAIS – C1',
                             '<LANGUAGE2>':'ITALIEN – B2',
                             '<LANGUAGE3>':'ESPAGNOL – B2'})
doc = addSkills(doc, ['python', 'matlab', 'keras', 'tensorflow', 'sklearn'])
doc = addPicture(doc, 'profile_picture.png')
doc.save("output.docx")