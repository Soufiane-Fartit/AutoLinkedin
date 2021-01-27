import argparse
import json
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from utils import getSkills, clusterSkills, merge_skills

def add_header_without_picture(doc, header):
    """ADD THE HEADER OF THE RESUME (NAME, ADRESS, EMAIL...)

    Args:
        doc ([docx document]): [description]
        header ([json]): json containing the name, adress, email ..

    Returns:
        [docx document]: the same docx document with the header added to it
    """
    for i, elem in enumerate(header) :
        if i == 0 :
            p = doc.add_paragraph(elem)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 0
            p.paragraph_format.space_before = Cm(0)
            p.paragraph_format.space_after = Cm(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            for run in p.runs :
                run.font.size = Pt(14)
        else :
            p = doc.add_paragraph(elem)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 0
            p.paragraph_format.space_before = Cm(0)
            p.paragraph_format.space_after = Cm(0.4)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    return doc

def add_header_with_picture(doc, header, picture_path):
    """ADD THE HEADER OF THE RESUME (NAME, ADRESS, EMAIL...)

    Args:
        doc ([docx document]): [description]
        header ([json]): json containing the name, adress, email ..

    Returns:
        [docx document]: the same docx document with the header added to it
    """

    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].width = Inches(4.7)
    table.rows[0].cells[1].width = Inches(2)
    table.rows[0].cells[2].width = Inches(1.7)

    # ADD THE PICTURE
    cell = table.cell(0, 2)
    cell._element.clear_content()
    p = cell.add_paragraph()
    p.add_run().add_picture(picture_path, width=Inches(1))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 0
    p.paragraph_format.space_before = Cm(0)
    p.paragraph_format.space_after = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    # ADD NAME
    cell = table.cell(0, 0)
    cell._element.clear_content()
    p = cell.add_paragraph(header['name'])
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 0
    p.paragraph_format.space_before = Cm(0)
    p.paragraph_format.space_after = Cm(0.6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    for run in p.runs :
        run.font.size = Pt(17)

    # ADD THE REST OF INFOS
    lines = [['city', 'phone', 'email'],
            ['linkedin', 'github']]
    for line in lines :
        p = cell.add_paragraph(' | '.join([header[x] for x in line]))
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 0
        p.paragraph_format.space_before = Cm(0)
        p.paragraph_format.space_after = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    return doc

def add_header(doc, header, picture_path=None):
    if picture_path :
        doc = add_header_with_picture(doc, header, picture_path)
    else:
        doc = add_header_without_picture(doc, header)
    return doc

def add_profil(doc, profil):
    """ADD FEW LINES ABOUT THE PERSON

    Args:
        doc ([type]): [description]
        profil ([type]): [description]

    Returns:
        [type]: [description]
    """
    doc = add_title(doc, "PROFIL")
    for elem in profil :
        p = doc.add_paragraph(elem)
        p.paragraph_format.line_spacing = 1

    return doc

def add_formation(doc, formation):
    """ADD EDUCATION SECTION TO THE RESUME

    Args:
        doc ([type]): [description]
        formation ([type]): [description]

    Returns:
        [type]: [description]
    """
    doc = add_title(doc, "FORMATION ACADEMIQUE")
    for form in formation :
        p = doc.add_paragraph(form['date']+' : '+form['diplome']+' - '+form['ecole'])
        p.paragraph_format.line_spacing = 0
        p.paragraph_format.space_before = Cm(0)
        p.paragraph_format.space_after = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    return doc

def add_experience(doc, experience):
    """ADD EXPERIENCE SECTION TO THE RESUME

    Args:
        doc ([type]): [description]
        experience ([type]): [description]

    Returns:
        [type]: [description]
    """
    doc = add_title(doc, "EXPERIENCE PROFESSIONNELLE")
    for exp in experience :
        # ADD JOB TITLE LINE
        p = doc.add_paragraph(exp['title'])
        p.paragraph_format.space_before = Cm(0.2)
        p.paragraph_format.space_after = Cm(0)
        p.paragraph_format.line_spacing = 1
        for run in p.runs:
            run.bold = True
        
        # ADD COMPANY NAME, ADRESS, DATES
        p = doc.add_paragraph(exp['company'] + ', '+ (9*"\t").join([exp['site'], exp['date']]))
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_before = Cm(0) 
        p.paragraph_format.space_after = Cm(0) 
        for run in p.runs:
            run.italic = True
        
        # ADD CONTENT
        if exp['text'] != "" :
            p = doc.add_paragraph(exp['text'])
            p.paragraph_format.line_spacing = 1
        
        # ADD BULLETPOINTS
        for task in exp['tasks'] :
            p = doc.add_paragraph(task, style='List Bullet')
            p.paragraph_format.line_spacing = 1
            p.paragraph_format.space_before = Cm(0)
            p.paragraph_format.space_after = Cm(0)
        
        doc.add_paragraph('Environnement : ' + ', '.join(exp['Environnement']))

    return doc

def add_skills(doc, skills, job_skills, n_cols=1):
    """ADD SKILLS SECTION TO THE RESUME

    Args:
        doc ([type]): [description]
        experience ([type]): [description]

    Returns:
        [type]: [description]
    """
    doc = add_title(doc, "COMPETENCES TECHNIQUES")

    #print(120*'-')
    #print(f'skills : {skills}')
    #print(f'job_skills : {job_skills}')

    if n_cols == 1 :
        default_skills = [item.lower() for sublist in skills.values() for item in sublist]
        missing_skills = [x.lower() for x in job_skills if x.lower() not in default_skills]
        skills_dict, skills_keys = getSkills('../../data/Input/Skillset-IT.xlsx')
        missing_skills_clustered = clusterSkills(missing_skills, skills_dict)
        merged_skills = merge_skills(missing_skills_clustered, skills)

        # SORT SKILLS BASED ON COLUMNS ORDER IN EXCEL FILE
        merged_skills_ = {}
        for key in skills_keys:
            try :
                merged_skills_[key] = merged_skills[key]
            except :
                pass
        merged_skills = merged_skills_
        #print(f'merged_skills : {merged_skills}')

        for k, v in merged_skills.items():
            if len(v) !=0 :
                p = doc.add_paragraph(k + ' : ' + ', '.join(v))
                p.paragraph_format.line_spacing = 1
                p.paragraph_format.space_before = Cm(0)
                p.paragraph_format.space_after = Cm(0)
    
    else :
        section = doc.add_section(WD_SECTION.CONTINUOUS)
        set_number_of_columns(section, n_cols)

        default_skills = [item.lower() for sublist in skills.values() for item in sublist]
        missing_skills = [x.lower() for x in job_skills if x.lower() not in default_skills]
        skills_dict, skills_keys = getSkills('../../data/Input/Skillset-IT.xlsx')
        missing_skills_clustered = clusterSkills(missing_skills, skills_dict)
        merged_skills = merge_skills(missing_skills_clustered, skills)
        
        # SORT SKILLS BASED ON COLUMNS ORDER IN EXCEL FILE
        merged_skills_ = {}
        for key in skills_keys:
            try :
                merged_skills_[key] = merged_skills[key]
            except :
                pass
        merged_skills = merged_skills_
        #print(f'merged_skills : {merged_skills}')

        for k, v in merged_skills.items():
            if len(v) !=0 :
                p = doc.add_paragraph(k + ' : ' + ', '.join(v))
                p.paragraph_format.line_spacing = 1
                p.paragraph_format.space_before = Cm(0)
                p.paragraph_format.space_after = Cm(0)
        
        section = doc.add_section(WD_SECTION.CONTINUOUS)
        set_number_of_columns(section, 1)

    return doc

def add_certif_projects(doc, certifs, projects, n_cols=1):
    """ADD CERTIFICATIONS AND PROJECTS SECTION TO THE RESUME

    Args:
        doc ([type]): [description]
        experience ([type]): [description]

    Returns:
        [type]: [description]
    """
    doc = add_title(doc, "CERTIFICATIONS ET PROJETS PERSONNELS")

    if n_cols == 1 :
        for certif in certifs:
            p = doc.add_paragraph(certif, style='List Bullet')
            p.paragraph_format.line_spacing = 1
        for project in projects:
            p = doc.add_paragraph(project, style='List Bullet')
            p.paragraph_format.line_spacing = 1
    else :
        section = doc.add_section(WD_SECTION.CONTINUOUS)
        set_number_of_columns(section, n_cols)
        
        for certif in certifs:
            p = doc.add_paragraph(certif, style='List Bullet')
            p.paragraph_format.line_spacing = 1
        for project in projects:
            p = doc.add_paragraph(project, style='List Bullet')
            p.paragraph_format.line_spacing = 1
        
        section = doc.add_section(WD_SECTION.CONTINUOUS)
        set_number_of_columns(section, 1)
    
    return doc

def add_langues(doc, langues):
    """ADD LANGUAGES SECTION TO THE RESUME

    Args:
        doc ([type]): [description]
        experience ([type]): [description]

    Returns:
        [type]: [description]
    """
    doc = add_title(doc, "LANGUES")
    
    p = doc.add_paragraph((5*' '+3*'\t').join([k+' : '+v for k,v in langues.items()]))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_before = Cm(0) 
    p.paragraph_format.space_after = Cm(0)

    return doc

def set_margin(doc, margin = 1.0):
    """SET THE MARGIN OF THE DOCX TO MAKE OT WIDER

    Args:
        doc ([type]): [description]
        margin (float, optional): [description]. Defaults to 1.0.

    Returns:
        [type]: [description]
    """
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin) 
    
    return doc

def first_child_found_in(parent, tagnames):
    """
    Return the first child of parent with tag in *tagnames*, or None if
    not found.
    """
    for tagname in tagnames:
        child = parent.find(qn(tagname))
        if child is not None:
            return child
    return None

def insert_element_before(parent, elm, successors):
    """
    Insert *elm* as child of *parent* before any existing child having
    tag name found in *successors*.
    """
    successor = first_child_found_in(parent, successors)
    if successor is not None:
        successor.addprevious(elm)
    else:
        parent.append(elm)
    return elm

def set_number_of_columns(section, cols):
    """ sets number of columns through xpath. """
    WNS_COLS_NUM = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"
    section._sectPr.xpath("./w:cols")[0].set(WNS_COLS_NUM, str(cols))

def add_title(doc, title):
    """
    ADD THE TITLE OF A SECTION
    """
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_before = Cm(0.5)
    pp.paragraph_format.space_after = Cm(0.1)
    run = pp.add_run(title)
    #run.bold = True
    """

    # ADD COLOR SHADING

    # Get the XML tag
    tag = run._r

    # Create XML element
    shd = OxmlElement('w:shd')

    # Add attributes to the element
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '90a4ae')

    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 255, 255)
    tag.rPr.append(shd)
    """
    p = pp._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    insert_element_before(pPr, pBdr, successors=(
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    ))
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

    return doc

def redact_cv(doc, header_, profil, formation, experience, skills, certifs, projects, langues, job_title, job_skills, picture_path=None, n_cols=1):
    """
    CREATE A RESUME FROM A JSON CONTAINING INFOS ABOUT A JOB

    Args:
        doc ([type]): [description]
        header ([type]): [description]
        profil ([type]): [description]
        formation ([type]): [description]
        experience ([type]): [description]
        skills ([type]): [description]
        certifs ([type]): [description]
        projects ([type]): [description]
        langues ([type]): [description]

    Returns:
        [docx document]: THE FULL RESUME GENERATED
    """

    # IF USING A PICTURE
    # WE USE TWO COLUMNS IN THE SKILLS AND CERTIFICATIONS SECTIONS
    # TO FIT THE RESUME IN ONE PAGE
    if picture_path and n_cols==1:
        n_cols = 2

    # ADD NAME, PHONE, ADRESS, EMAIL ...

    # PICTURE PATH IS NOT MANDATORY
    doc = add_header(doc, header_, picture_path)

    # ADD A FEW LINES ABOUT ME
    doc = add_profil(doc, [x.replace('<job_title>', job_title) for x in profil])

    # ADD PRO EXP SECTION
    doc = add_experience(doc, experience)

    # ADD EDUCATION SECTION
    doc = add_formation(doc, formation)

    # ADD SKILLS SECTION
    doc = add_skills(doc, skills, job_skills, n_cols=n_cols)

    # ADD CERTIFS PERS PROJECTS
    doc = add_certif_projects(doc, certifs, projects, n_cols=n_cols)

    # ADD SPOKE LANGUAGES
    doc = add_langues(doc, langues)

    # SET MARGINS
    doc = set_margin(doc)

    return doc

def main():
    # GET ARGUMENTS
    parser = argparse.ArgumentParser()
    parser.add_argument("-c", "--columns", type=int, help="numnber of columns for skills and certificates sections")
    parser.add_argument("-p", "--picture", type=str, help="path to the picture")
    args = parser.parse_args()
    n_cols = args.columns
    picture_path = args.picture

    assert not(n_cols==1 and picture_path), "you can't set columns=1 when using a picture"

    if n_cols==None and picture_path==None:
        n_cols = 1

    with open('../../data/Input/resume_data.json', 'r') as f:
        d = json.load(f)
        header, header_with_picture_temporary_element, profil, formation, experience, skills, certifs, projects, langues = d['header'], d['header_with_picture_temporary_element'], d['profil'], d['formation'], d['experience'], d['skills'], d['certifs'], d['projects'], d['langues']

    # LOAD PARSED DATA
    with open('../../data/ScrapedJobsData/jobsFoundParsed.json', 'r') as jobs_file:
        jobs = json.load(jobs_file)

    # ITERATE OVER JOBS AND GENERATE CUSTOM RESUMES
    for i, job in enumerate(jobs) :
        #print(f"---------------- {i} ----------------")
        job_skills = job['skills_found']
        job_title = job['title']
        
        jobs[i]['id_code'] = i
        jobs[i]['filename'] = 'CV_'+str(i)+'.docx'

        document = Document()

        
        if picture_path :
            # WITH PICTURE
            document = redact_cv(document, header_with_picture_temporary_element, profil, formation, experience, skills, certifs, projects, langues, job_title, job_skills, picture_path=picture_path, n_cols=n_cols)
        else :
            # WITHOUT PICTURE
            document = redact_cv(document, header, profil, formation, experience, skills, certifs, projects, langues, job_title, job_skills, n_cols=n_cols)
        
        document.save('../../data/Resumes/CV_'+str(i)+'.docx')

    # SAVE FILE NAMES OF THE RESUMES IN THE JSON FILE
    with open('../../data/ScrapedJobsData/jobsFoundParsed.json', 'w') as fout:
        jobs = [job for job in jobs if 'skills_found' in job]
        json.dump(jobs, fout, ensure_ascii=False, indent = 4)

if __name__ == '__main__':
    main()